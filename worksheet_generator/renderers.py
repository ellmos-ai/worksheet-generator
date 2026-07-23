"""Renderer-Adapter: wandeln ein Arbeitsblatt-JSON (schema.py) in
Ausgabeformate um.

Kern (immer verfuegbar, keine Zusatzabhaengigkeiten): Markdown, HTML.
Optional: DOCX (python-docx) -- faellt sauber mit RendererUnavailable aus,
wenn die Bibliothek nicht installiert ist.

PDF wird bewusst NICHT selbst erzeugt: HTML -> PDF je nach lokalem
Werkzeugbestand extern erledigen (z.B. `cc_md_to_pdf`/`fc_md_to_pdf`
[ellmos CodeCommander/FileCommander MCP], ein Browser-Druckdialog, oder
pandoc). Details: SKILL.md, Abschnitt "Rendern".

Weitere optionale Adapter (PowerPoint, Canva) delegieren an externe
Design-Skills/Connectoren und liegen bewusst AUSSERHALB dieses Kern-Moduls
(siehe SKILL.md, Abschnitt "Design-Delegation").
"""
from __future__ import annotations

from pathlib import Path
from typing import Any

_KIND_LABELS = {
    "luecke": "Luecke",
    "zuordnung": "Zuordnung",
    "rechnung": "Rechenaufgabe",
    "frage": "Frage",
    "freitext": "Freitext",
}


class RendererUnavailable(RuntimeError):
    """Ausgeloest, wenn ein optionaler Renderer (z.B. docx) nicht installiert ist."""


def _meta_bits(goal: dict[str, Any]) -> list[str]:
    bits = []
    if goal.get("icf_codes"):
        titles = goal.get("icf_titles", {})
        codes_fmt = ", ".join(
            f"{c} ({titles[c]})" if c in titles else c for c in goal["icf_codes"]
        )
        bits.append(f"ICF: {codes_fmt}")
    if goal.get("niveau"):
        bits.append(f"Niveau: {goal['niveau']}")
    if goal.get("alter"):
        bits.append(f"Alter: {goal['alter']}")
    return bits


def to_markdown(worksheet: dict[str, Any]) -> str:
    """Rendert ein Arbeitsblatt-JSON nach Markdown."""
    meta = worksheet["meta"]
    goal = meta.get("goal", {})
    lines = [f"# {meta['title']}", ""]
    bits = _meta_bits(goal)
    if bits:
        lines.append(f"*{' | '.join(bits)}*")
        lines.append("")

    for section in worksheet.get("sections", []):
        lines.append(f"## {section['title']}")
        if section.get("content"):
            lines.append(section["content"])
        lines.append("")
        for idx, item in enumerate(section.get("items", []), start=1):
            label = _KIND_LABELS.get(item["kind"], item["kind"])
            lines.append(f"{idx}. **[{label}]** {item['prompt']}")
            if item.get("options"):
                lines.append("   Optionen: " + ", ".join(item["options"]))
            if item.get("hint"):
                lines.append(f"   Hinweis: {item['hint']}")
            if item.get("answer_space", True):
                lines.append("   " + "_" * 30)
        lines.append("")
    return "\n".join(lines).rstrip() + "\n"


def to_html(worksheet: dict[str, Any]) -> str:
    """Rendert ein Arbeitsblatt-JSON nach eigenstaendigem HTML (druckbar)."""
    import html as _html

    meta = worksheet["meta"]
    goal = meta.get("goal", {})
    parts = [
        "<!doctype html>",
        '<html lang="de"><head><meta charset="utf-8">',
        f"<title>{_html.escape(meta['title'])}</title>",
        "<style>body{font-family:sans-serif;max-width:700px;margin:2em auto;line-height:1.5}"
        "h1{margin-bottom:0.2em}.meta{color:#555;font-style:italic}"
        "ol{padding-left:1.2em}.answer-space{border-bottom:1px solid #999;"
        "display:block;height:1.4em;margin:0.3em 0}</style>",
        "</head><body>",
        f"<h1>{_html.escape(meta['title'])}</h1>",
    ]
    bits = _meta_bits(goal)
    if bits:
        parts.append(f'<p class="meta">{_html.escape(" | ".join(bits))}</p>')

    for section in worksheet.get("sections", []):
        parts.append(f"<h2>{_html.escape(section['title'])}</h2>")
        if section.get("content"):
            parts.append(f"<p>{_html.escape(section['content'])}</p>")
        items = section.get("items", [])
        if items:
            parts.append("<ol>")
            for item in items:
                label = _KIND_LABELS.get(item["kind"], item["kind"])
                parts.append(
                    f"<li><strong>[{_html.escape(label)}]</strong> "
                    f"{_html.escape(item['prompt'])}"
                )
                if item.get("options"):
                    parts.append("<br>Optionen: " + _html.escape(", ".join(item["options"])))
                if item.get("hint"):
                    parts.append(f"<br>Hinweis: {_html.escape(item['hint'])}")
                if item.get("answer_space", True):
                    parts.append('<span class="answer-space"></span>')
                parts.append("</li>")
            parts.append("</ol>")
    parts.append("</body></html>")
    return "\n".join(parts)


def to_docx(worksheet: dict[str, Any], out_path: str | Path) -> Path:
    """Schreibt ein schlichtes DOCX aus dem Arbeitsblatt-JSON.

    Erfordert die optionale Abhaengigkeit python-docx; wirft
    RendererUnavailable, wenn sie nicht installiert ist.
    """
    try:
        import docx
    except ImportError as exc:
        raise RendererUnavailable(
            "python-docx nicht installiert -- 'pip install python-docx' fuer DOCX-Export."
        ) from exc

    meta = worksheet["meta"]
    goal = meta.get("goal", {})
    document = docx.Document()
    document.add_heading(meta["title"], level=1)

    bits = _meta_bits(goal)
    if bits:
        p = document.add_paragraph(" | ".join(bits))
        if p.runs:
            p.runs[0].italic = True

    for section in worksheet.get("sections", []):
        document.add_heading(section["title"], level=2)
        if section.get("content"):
            document.add_paragraph(section["content"])
        for idx, item in enumerate(section.get("items", []), start=1):
            label = _KIND_LABELS.get(item["kind"], item["kind"])
            document.add_paragraph(f"{idx}. [{label}] {item['prompt']}")
            if item.get("options"):
                document.add_paragraph("   Optionen: " + ", ".join(item["options"]))
            if item.get("hint"):
                document.add_paragraph(f"   Hinweis: {item['hint']}")
            if item.get("answer_space", True):
                document.add_paragraph("   " + "_" * 30)

    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(out))
    return out


# Kern-Renderer, die ohne Zusatzabhaengigkeiten laufen (fuer cli.py/status).
RENDERERS = {"md": to_markdown, "html": to_html}
